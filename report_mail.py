
# libraries to be imported
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import pdfkit
from docx import Document


def mail_send(data,receiver):
    document = Document()

    document.add_heading('Personality Traits', 0)

    results = data.split("#")
    p = document.add_paragraph(results[0])
    document.add_heading('Career Suggestions', 0)
    p1 = document.add_paragraph(results[1]+"\n"+results[2]+"\n"+results[3]+"\n"+results[7]+"\n"+results[5]+"\n"+results[6]+"\n");

    document.save('Traits.docx')



    fromaddr = "grapholearning1@gmail.com"
    toaddr = receiver

    # instance of MIMEMultipart
    msg = MIMEMultipart()

    # storing the senders email address
    msg['From'] = fromaddr

    # storing the receivers email address
    msg['To'] = toaddr

    # storing the subject
    msg['Subject'] = "Personality Traits"

    # string to store the body of the mail
    body = """" 
            <html>
            <body>
                  <h1> Grapho-Learning </h1> 
            </body>
            </html>
        """

    # attach the body with the msg instance
    msg.attach(MIMEText(body, 'html'))

    # open the file to be sent
    filename = "Traits.docx"
    attachment = open("Traits.docx", "rb")

    # instance of MIMEBase and named as p
    p = MIMEBase('application', 'octet-stream')

    # To change the payload into encoded form
    p.set_payload((attachment).read())

    # encode into base64
    encoders.encode_base64(p)

    p.add_header('Content-Disposition', "attachment; filename= %s" % filename)

    # attach the instance 'p' to instance 'msg'
    msg.attach(p)

    # creates SMTP session
    s = smtplib.SMTP('smtp.gmail.com', 587)

    # start TLS for security
    s.starttls()

    # Authentication
    s.login(fromaddr, "grapho123")

    # Converts the Multipart msg into a string
    text = msg.as_string()

    # sending the mail
    s.sendmail(fromaddr, toaddr, text)

    # terminating the session
    s.quit()
    print("data sent")
    return True

#mail_send('Hello How are you')